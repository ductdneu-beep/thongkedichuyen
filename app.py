import streamlit as st
import pandas as pd
import math
import json
import os
import io
import re
import urllib.request
from docx import Document
from docx.shared import Pt, RGBColor
from docx.enum.text import WD_ALIGN_PARAGRAPH
from reportlab.lib.pagesizes import A4
from reportlab.lib import colors
from reportlab.platypus import SimpleDocTemplate, Paragraph, Spacer, Table, TableStyle
from reportlab.lib.styles import getSampleStyleSheet, ParagraphStyle
from reportlab.pdfbase import pdfmetrics
from reportlab.pdfbase.ttfonts import TTFont

st.set_page_config(
    page_title="Phương án di chuyển",
    page_icon="✈️",
    layout="wide"
)

# -----------------------------------------------------------------------------
# CẤU HÌNH FONT UNICODE (TIMES NEW ROMAN / SERIF) TỰ ĐỘNG CHO MỌI HỆ ĐIỀU HÀNH
# -----------------------------------------------------------------------------
@st.cache_resource
def setup_pdf_fonts():
    font_name_reg = "VN-Times"
    font_name_bold = "VN-Times-Bold"
    
    # Danh sách kiểm tra font có sẵn trên macOS / Linux
    local_reg_paths = [
        "/System/Library/Fonts/Supplemental/Times New Roman.ttf",
        "/Library/Fonts/Times New Roman.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSerif.ttf"
    ]
    local_bold_paths = [
        "/System/Library/Fonts/Supplemental/Times New Roman Bold.ttf",
        "/Library/Fonts/Times New Roman Bold.ttf",
        "/usr/share/fonts/truetype/msttcorefonts/Times_New_Roman_Bold.ttf",
        "/usr/share/fonts/truetype/dejavu/DejaVuSerif-Bold.ttf",
        "/usr/share/fonts/truetype/freefont/FreeSerifBold.ttf"
    ]
    
    reg_loaded = False
    bold_loaded = False
    
    for p in local_reg_paths:
        if os.path.exists(p):
            try:
                pdfmetrics.registerFont(TTFont(font_name_reg, p))
                reg_loaded = True
                break
            except Exception:
                pass
                
    for p in local_bold_paths:
        if os.path.exists(p):
            try:
                pdfmetrics.registerFont(TTFont(font_name_bold, p))
                bold_loaded = True
                break
            except Exception:
                pass
                
    # Nếu đang chạy trên Streamlit Cloud (chưa có font local), tự động tải font serif Unicode về cache
    if not reg_loaded or not bold_loaded:
        cache_dir = os.path.expanduser("~/.fonts_cache")
        os.makedirs(cache_dir, exist_ok=True)
        reg_file = os.path.join(cache_dir, "LiberationSerif-Regular.ttf")
        bold_file = os.path.join(cache_dir, "LiberationSerif-Bold.ttf")
        
        try:
            if not os.path.exists(reg_file):
                urllib.request.urlretrieve(
                    "https://github.com/google/fonts/raw/main/apache/robotoserif/RobotoSerif%5BGRAD%2Copsz%2Cwdth%2Cwght%5D.ttf",
                    reg_file
                )
            pdfmetrics.registerFont(TTFont(font_name_reg, reg_file))
            reg_loaded = True
        except Exception:
            pass

        try:
            if not os.path.exists(bold_file):
                bold_file = reg_file # dùng chung nếu cần
            pdfmetrics.registerFont(TTFont(font_name_bold, bold_file))
            bold_loaded = True
        except Exception:
            pass

    return font_name_reg if reg_loaded else "Helvetica", font_name_bold if bold_loaded else "Helvetica-Bold"

PDF_FONT_REG, PDF_FONT_BOLD = setup_pdf_fonts()

SAVED_DB_FILE = os.path.expanduser("~/.saved_plans_db.json")

def load_saved_db():
    if os.path.exists(SAVED_DB_FILE):
        try:
            with open(SAVED_DB_FILE, "r", encoding="utf-8") as f:
                return json.load(f)
        except Exception:
            return {}
    return {}

def save_saved_db(db):
    try:
        with open(SAVED_DB_FILE, "w", encoding="utf-8") as f:
            json.dump(db, f, ensure_ascii=False, indent=2)
    except Exception as e:
        st.error(f"Không thể lưu: {e}")

# -----------------------------------------------------------------------------
# 1. TỌA ĐỘ 63 TỈNH THÀNH
# -----------------------------------------------------------------------------
PROVINCES = {
    "An Giang": (10.5361, 105.1259), "Bà Rịa - Vũng Tàu": (10.5417, 107.2429), "Bắc Giang": (21.2731, 106.1946),
    "Bắc Kạn": (22.1470, 105.8348), "Bạc Liêu": (9.2941, 105.7244), "Bắc Ninh": (21.1861, 106.0763),
    "Bến Tre": (10.2432, 106.3752), "Bình Định": (13.7820, 109.2194), "Bình Dương": (11.1604, 106.6519),
    "Bình Phước": (11.7512, 106.9184), "Bình Thuận": (11.0904, 108.0721), "Cà Mau": (9.1769, 105.1524),
    "Cần Thơ": (10.0452, 105.7469), "Cao Bằng": (22.6657, 105.9752), "Đà Nẵng": (16.0544, 108.2022),
    "Đắk Lắk": (12.6667, 108.0500), "Đắk Nông": (12.0042, 107.6875), "Điện Biên": (21.3842, 103.0238),
    "Đồng Nai": (10.9574, 106.8427), "Đồng Tháp": (10.4938, 105.6373), "Gia Lai": (13.9833, 108.0000),
    "Hà Giang": (22.8233, 104.9839), "Hà Nam": (20.5839, 105.9242), "Hà Nội": (21.0285, 105.8542),
    "Hà Tĩnh": (18.3428, 105.9056), "Hải Dương": (20.9381, 106.3194), "Hải Phòng": (20.8449, 106.6881),
    "Hậu Giang": (9.7839, 105.4702), "Hòa Bình": (20.8172, 105.3378), "Hưng Yên": (20.6464, 106.0511),
    "Khánh Hòa": (12.2388, 109.1967), "Kiên Giang": (10.0125, 105.0809), "Kon Tum": (14.3500, 108.0000),
    "Lai Châu": (22.3964, 103.4581), "Lâm Đồng": (11.9404, 108.4583), "Lạng Sơn": (21.8472, 106.7583),
    "Lào Cai": (22.4856, 103.9707), "Long An": (10.5361, 106.4025), "Nam Định": (20.4286, 106.1683),
    "Nghệ An": (18.6734, 105.6813), "Ninh Bình": (20.2539, 105.9750), "Ninh Thuận": (11.5653, 108.9886),
    "Phú Thọ": (21.3228, 105.2281), "Phú Yên": (13.0882, 109.3113), "Quảng Bình": (17.4761, 106.5983),
    "Quảng Nam": (15.5681, 108.1272), "Quảng Ngãi": (15.1206, 108.7922), "Quảng Ninh": (21.0069, 107.2925),
    "Quảng Trị": (16.7411, 107.1856), "Sóc Trăng": (9.6033, 105.9800), "Sơn La": (21.3256, 103.9189),
    "Tây Ninh": (11.3125, 106.0989), "Thái Bình": (20.4464, 106.3364), "Thái Nguyên": (21.5928, 105.8442),
    "Thanh Hóa": (19.8067, 105.7853), "Thừa Thiên Huế": (16.4637, 107.5909), "Tiền Giang": (10.3602, 106.3614),
    "TP.HCM": (10.8231, 106.6297), "Trà Vinh": (9.9347, 106.3444), "Tuyên Quang": (21.8233, 105.2167),
    "Vĩnh Long": (10.2537, 105.9722), "Vĩnh Phúc": (21.3089, 105.6047), "Yên Bái": (21.7167, 104.9000)
}

DEFAULT_SPECS = [
    {"default_name": "Máy bay", "speed": 550, "cost_per_km": 2100, "base_cost": 450000, "fixed_time": 2.5, "min_dist": 350},
    {"default_name": "Tàu hỏa", "speed": 60, "cost_per_km": 850, "base_cost": 80000, "fixed_time": 1.0, "min_dist": 120},
    {"default_name": "Xe khách / Limousine", "speed": 60, "cost_per_km": 650, "base_cost": 40000, "fixed_time": 0.5, "min_dist": 0},
    {"default_name": "Ô tô riêng / Thuê xe", "speed": 65, "cost_per_km": 1000, "base_cost": 0, "fixed_time": 0.2, "min_dist": 0}
]

def calculate_haversine(origin, destination):
    if origin == destination: return 0
    lat1, lon1 = PROVINCES[origin]
    lat2, lon2 = PROVINCES[destination]
    R = 6371
    dlat = math.radians(lat2 - lat1)
    dlon = math.radians(lon2 - lon1)
    a = math.sin(dlat/2)**2 + math.cos(math.radians(lat1)) * math.cos(math.radians(lat2)) * math.sin(dlon/2)**2
    return round(R * (2 * math.atan2(math.sqrt(a), math.sqrt(1-a))) * 1.3)

def format_mode_transport_label(mode, flight_code):
    m_lower = str(mode).lower()
    code_str = str(flight_code).strip() if flight_code else ""
    if any(k in m_lower for k in ["bay", "flight", "plane", "vietnam airlines", "vietjet", "bamboo"]):
        return f"{mode} (Mã chuyến bay: {code_str})" if code_str else mode
    elif any(k in m_lower for k in ["thủy", "cano", "ca nô", "tàu", "hỏa", "thuyền", "phà", "railway", "train", "boat"]):
        return f"{mode} (Mã chuyến đi: {code_str})" if code_str else mode
    else:
        return mode

def generate_itinerary_steps(origin, destination, mode, flight_code, dist):
    if dist == 0:
        return "Tập trung tại chỗ (Không cần di chuyển xa)"
    m_lower = str(mode).lower()
    code_str = str(flight_code).strip() if flight_code else ""
    if any(k in m_lower for k in ["bay", "flight", "plane", "vietnam airlines", "vietjet", "bamboo"]):
        step2 = f"Bước 2: Di chuyển bằng {mode}" + (f" (Mã chuyến bay: {code_str})" if code_str else "") + f" về {destination} ({dist} km)."
    elif any(k in m_lower for k in ["thủy", "cano", "ca nô", "tàu", "hỏa", "thuyền", "phà", "railway", "train", "boat"]):
        step2 = f"Bước 2: Di chuyển bằng {mode}" + (f" (Mã chuyến đi: {code_str})" if code_str else "") + f" về {destination} ({dist} km)."
    else:
        step2 = f"Bước 2: Di chuyển bằng {mode} về {destination} ({dist} km)."
    return f"Bước 1: Di chuyển từ {origin} ra điểm khởi hành (~1h).\n{step2}\nBước 3: Đón xe về điểm hội quân C."

def clean_pdf_text(text):
    if not text: return ""
    text = str(text)
    text = text.replace("➔", "->").replace("➜", "->").replace("➡", "->").replace("→", "->")
    text = text.replace("—", " - ").replace("–", " - ").replace("•", "-")
    emoji_pattern = re.compile(
        "[\U00010000-\U0010ffff\U00002600-\U000027BF\U0001f300-\U0001f64f\U0001f680-\U0001f6ff\U0001f1e0-\U0001f1ff]+",
        flags=re.UNICODE
    )
    return emoji_pattern.sub('', text).strip()

def generate_docx(results, plan_title, destination, total_cost, max_dur):
    doc = Document()
    style = doc.styles['Normal']
    font = style.font
    font.name = 'Times New Roman'
    font.size = Pt(12)
    
    title = doc.add_heading("BÁO CÁO PHƯƠNG ÁN DI CHUYỂN", level=0)
    title.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for run in title.runs:
        run.font.name = 'Times New Roman'
        run.font.size = Pt(18)
        run.font.bold = True
        run.font.color.rgb = RGBColor(0, 0, 0)
    
    p_sub = doc.add_paragraph(f"Điểm đến hội quân: {destination} | Tiêu chí: {plan_title}")
    p_sub.alignment = WD_ALIGN_PARAGRAPH.CENTER
    for r in p_sub.runs: r.font.name = 'Times New Roman'
    
    p_sum = doc.add_paragraph(f"- Tổng chi phí toàn đoàn: {total_cost:,} VNĐ\n- Thời gian toàn đoàn đến đủ: {max_dur} Giờ\n- Tổng số lượng khách: {sum(r['people'] for r in results)} Người")
    for r in p_sum.runs: r.font.name = 'Times New Roman'
    
    h1 = doc.add_heading("1. Bảng Tổng Hợp Chi Phí & Phương Tiện", level=1)
    for r in h1.runs: r.font.name = 'Times New Roman'; r.font.bold = True; r.font.color.rgb = RGBColor(0, 51, 102)
    
    table = doc.add_table(rows=1, cols=7)
    table.style = 'Light Grid Accent 1'
    hdr_cells = table.rows[0].cells
    headers = ["Đoàn", "Xuất phát", "Số người", "Phương tiện", "Quãng đường", "Thời gian", "Tổng chi phí"]
    for i, h in enumerate(headers):
        hdr_cells[i].text = h
        for run in hdr_cells[i].paragraphs[0].runs:
            run.font.name = 'Times New Roman'
            run.font.bold = True
        
    for res in results:
        row_cells = table.add_row().cells
        mode_str = format_mode_transport_label(res['mode'], res['flight_code'])
        vals = [str(res['name']), str(res['origin']), str(res['people']), str(mode_str), f"{res['distance']} km", f"{res['duration']} h", f"{res['total_cost']:,} VNĐ"]
        for idx, val in enumerate(vals):
            row_cells[idx].text = val
            for run in row_cells[idx].paragraphs[0].runs:
                run.font.name = 'Times New Roman'
        
    doc.add_paragraph("")
    h2 = doc.add_heading("2. Lộ Trình Di Chuyển Chi Tiết Từng Đoàn", level=1)
    for r in h2.runs: r.font.name = 'Times New Roman'; r.font.bold = True; r.font.color.rgb = RGBColor(0, 51, 102)
    
    for res in results:
        head_d = doc.add_heading(f"- {res['name']} ({res['origin']} -> {destination})", level=2)
        for r in head_d.runs: r.font.name = 'Times New Roman'; r.font.bold = True
        
        mode_str = format_mode_transport_label(res['mode'], res['flight_code'])
        p_d = doc.add_paragraph(
            f"+ Số người: {res['people']} khách | Quãng đường: {res['distance']} km\n"
            f"+ Phương tiện: {mode_str}\n"
            f"+ Chi phí / người: {res['cost_per_person']:,} VNĐ | Tổng tiền: {res['total_cost']:,} VNĐ\n"
            f"+ Chi tiết từng bước:\n{res['steps']}\n"
        )
        for r in p_d.runs: r.font.name = 'Times New Roman'
        
    bio = io.BytesIO()
    doc.save(bio)
    bio.seek(0)
    return bio

def generate_pdf(results, plan_title, destination, total_cost, max_dur):
    bio = io.BytesIO()
    doc = SimpleDocTemplate(bio, pagesize=A4, rightMargin=30, leftMargin=30, topMargin=30, bottomMargin=30)
    styles = getSampleStyleSheet()
    
    title_style = ParagraphStyle('TitleStyle', fontName=PDF_FONT_BOLD, fontSize=16, alignment=1, spaceAfter=8, textColor=colors.HexColor("#111827"))
    sub_style = ParagraphStyle('SubStyle', fontName=PDF_FONT_REG, fontSize=10, alignment=1, spaceAfter=12, textColor=colors.HexColor("#4B5563"))
    norm_style = ParagraphStyle('NormStyle', fontName=PDF_FONT_REG, fontSize=10, leading=14, textColor=colors.black)
    head_style = ParagraphStyle('HeadStyle', fontName=PDF_FONT_BOLD, fontSize=12, spaceBefore=12, spaceAfter=8, textColor=colors.HexColor("#1E3A8A"))
    cell_style = ParagraphStyle('CellStyle', fontName=PDF_FONT_REG, fontSize=9, alignment=1, leading=12)
    cell_bold = ParagraphStyle('CellBold', fontName=PDF_FONT_BOLD, fontSize=9, alignment=1, leading=12)
    
    story = []
    story.append(Paragraph("BÁO CÁO PHƯƠNG ÁN DI CHUYỂN", title_style))
    story.append(Paragraph(f"Điểm đến hội quân: <b>{destination}</b> | Tiêu chí: <b>{plan_title}</b>", sub_style))
    story.append(Paragraph(f"- <b>Tổng chi phí:</b> {total_cost:,} VNĐ | <b>Thời gian đến đủ:</b> {max_dur} Giờ | <b>Tổng khách:</b> {sum(r['people'] for r in results)} Người", norm_style))
    story.append(Spacer(1, 10))
    
    headers = [Paragraph("Đoàn", cell_bold), Paragraph("Xuất phát", cell_bold), Paragraph("Số người", cell_bold), Paragraph("Phương tiện", cell_bold), Paragraph("Quãng đường", cell_bold), Paragraph("Thời gian", cell_bold), Paragraph("Tổng chi phí", cell_bold)]
    data = [headers]
    
    for r in results:
        mode_label = format_mode_transport_label(r['mode'], r['flight_code'])
        data.append([
            Paragraph(clean_pdf_text(r['name']), cell_style),
            Paragraph(clean_pdf_text(r['origin']), cell_style),
            Paragraph(str(r['people']), cell_style),
            Paragraph(clean_pdf_text(mode_label), cell_style),
            Paragraph(f"{r['distance']} km", cell_style),
            Paragraph(f"{r['duration']} h", cell_style),
            Paragraph(f"<b>{r['total_cost']:,}</b>", cell_style)
        ])
        
    table = Table(data, colWidths=[85, 75, 45, 115, 65, 55, 95])
    table.setStyle(TableStyle([
        ('BACKGROUND', (0, 0), (-1, 0), colors.HexColor("#F3F4F6")),
        ('ALIGN', (0, 0), (-1, -1), 'CENTER'),
        ('VALIGN', (0, 0), (-1, -1), 'MIDDLE'),
        ('BOTTOMPADDING', (0, 0), (-1, -1), 5),
        ('TOPPADDING', (0, 0), (-1, -1), 5),
        ('GRID', (0, 0), (-1, -1), 0.5, colors.HexColor("#D1D5DB")),
    ]))
    story.append(table)
    story.append(Spacer(1, 12))
    
    story.append(Paragraph("Chi Tiết Lộ Trình Hành Trình:", head_style))
    for r in results:
        mode_label = format_mode_transport_label(r['mode'], r['flight_code'])
        d_name = clean_pdf_text(r['name'])
        d_origin = clean_pdf_text(r['origin'])
        d_dest = clean_pdf_text(destination)
        d_mode = clean_pdf_text(mode_label)
        
        story.append(Paragraph(f"<b>- {d_name} ({d_origin} -> {d_dest}):</b> {d_mode} - <b>{r['total_cost']:,} VNĐ</b>", norm_style))
        clean_steps = clean_pdf_text(r['steps']).replace('\n', '<br/>')
        story.append(Paragraph(f"<font color='#374151'>{clean_steps}</font>", norm_style))
        story.append(Spacer(1, 6))
        
    doc.build(story)
    bio.seek(0)
    return bio

if "custom_configs" not in st.session_state: st.session_state.custom_configs = {}
if "override_mode" not in st.session_state: st.session_state.override_mode = {}
if "custom_itinerary" not in st.session_state: st.session_state.custom_itinerary = {}
if "mode_names_state" not in st.session_state: st.session_state.mode_names_state = {}
if "flight_codes" not in st.session_state: st.session_state.flight_codes = {}
if "groups" not in st.session_state:
    st.session_state.groups = [
        {"name": "Đoàn Miền Bắc", "origin": "Hà Nội", "people": 5},
        {"name": "Đoàn Tây Nguyên", "origin": "Đắk Lắk", "people": 4},
        {"name": "Đoàn Miền Tây", "origin": "Cần Thơ", "people": 8}
    ]

province_list = sorted(list(PROVINCES.keys()))

# -----------------------------------------------------------------------------
# 2. THANH BÊN (SIDEBAR)
# -----------------------------------------------------------------------------
with st.sidebar:
    st.header("⚙️ Cấu hình đề bài")
    destination = st.selectbox("📍 Điểm đến chung (C):", province_list, index=province_list.index("Đà Nẵng"))
    st.markdown("---")
    
    st.subheader("👥 Danh sách các đoàn")
    with st.expander("➕ Thêm đoàn mới", expanded=False):
        new_name = st.text_input("Tên đoàn:", value=f"Đoàn {len(st.session_state.groups)+1}")
        new_origin = st.selectbox("Tỉnh xuất phát:", province_list)
        new_people = st.number_input("Số người:", min_value=1, value=4)
        if st.button("Xác nhận thêm"):
            st.session_state.groups.append({"name": new_name, "origin": new_origin, "people": new_people})
            st.rerun()

    for idx, g in enumerate(st.session_state.groups):
        col1, col2 = st.columns([3, 1])
        col1.write(f"• **{g['name']}**: {g['people']} người ({g['origin']})")
        if col2.button("❌", key=f"del_{idx}"):
            st.session_state.groups.pop(idx)
            st.rerun()

    st.markdown("---")
    st.subheader("💾 Quản lý Lưu Trữ Dữ Liệu")
    saved_db = load_saved_db()
    
    plan_save_name = st.text_input("Tên bản lưu:", value="Phương án đoàn 1", placeholder="Nhập tên để dễ nhớ...")
    if st.button("💾 Lưu nhanh bản này vào phần mềm"):
        current_data = {
            "destination": destination,
            "groups": st.session_state.groups,
            "custom_configs": st.session_state.custom_configs,
            "override_mode": st.session_state.override_mode,
            "custom_itinerary": st.session_state.custom_itinerary,
            "mode_names_state": st.session_state.mode_names_state,
            "flight_codes": st.session_state.flight_codes
        }
        saved_db[plan_save_name] = current_data
        save_saved_db(saved_db)
        st.success(f"✅ Đã lưu '{plan_save_name}'!")
        st.rerun()
    
    if saved_db:
        st.markdown("**📂 Danh sách bản đã lưu:**")
        selected_plan_name = st.selectbox("Chọn bản muốn mở lại:", list(saved_db.keys()))
        col_load, col_del = st.columns(2)
        if col_load.button("🔄 Mở bản này"):
            plan_data = saved_db[selected_plan_name]
            st.session_state.groups = plan_data.get("groups", st.session_state.groups)
            st.session_state.custom_configs = plan_data.get("custom_configs", {})
            st.session_state.override_mode = plan_data.get("override_mode", {})
            st.session_state.custom_itinerary = plan_data.get("custom_itinerary", {})
            st.session_state.mode_names_state = plan_data.get("mode_names_state", {})
            st.session_state.flight_codes = plan_data.get("flight_codes", {})
            st.success(f"Đã mở: {selected_plan_name}")
            st.rerun()
        if col_del.button("🗑️ Xóa bản này"):
            del saved_db[selected_plan_name]
            save_saved_db(saved_db)
            st.warning(f"Đã xóa: {selected_plan_name}")
            st.rerun()
            
    st.markdown("---")
    custom_file_name = st.text_input("Tên file .json khi tải về:", value="du_lieu_doan")
    clean_filename = f"{custom_file_name.strip()}.json" if not custom_file_name.endswith(".json") else custom_file_name.strip()
    export_payload = {
        "destination": destination,
        "groups": st.session_state.groups,
        "custom_configs": st.session_state.custom_configs,
        "override_mode": st.session_state.override_mode,
        "custom_itinerary": st.session_state.custom_itinerary,
        "mode_names_state": st.session_state.mode_names_state,
        "flight_codes": st.session_state.flight_codes
    }
    st.download_button(label=f"📥 Tải file .json", data=json.dumps(export_payload, ensure_ascii=False, indent=2), file_name=clean_filename, mime="application/json")
    uploaded_file = st.file_uploader("📤 Mở file từ máy tính (.json):", type=["json"])
    if uploaded_file is not None and st.button("🔄 Nạp dữ liệu từ file"):
        try:
            loaded_data = json.load(uploaded_file)
            st.session_state.groups = loaded_data.get("groups", st.session_state.groups)
            st.session_state.custom_configs = loaded_data.get("custom_configs", {})
            st.session_state.override_mode = loaded_data.get("override_mode", {})
            st.session_state.custom_itinerary = loaded_data.get("custom_itinerary", {})
            st.session_state.mode_names_state = loaded_data.get("mode_names_state", {})
            st.session_state.flight_codes = loaded_data.get("flight_codes", {})
            st.success("✅ Đã nạp thành công!")
            st.rerun()
        except Exception as e: st.error(f"Lỗi: {e}")

st.title("🗺️ Phương án di chuyển")
st.caption("Tự động đề xuất phương tiện & cho phép tùy chỉnh giá vé, số hiệu chuyến bay, lộ trình và xuất báo cáo Word / PDF.")

# -----------------------------------------------------------------------------
# 3. BẢNG GIÁ
# -----------------------------------------------------------------------------
st.subheader("📝 Bảng giá")
st.caption("Bạn có thể nhập SỐ HIỆU CHUYẾN BAY/CHUYẾN ĐI, đổi tên phương tiện hoặc chọn lại phương tiện bên dưới:")

default_keys = ["m1", "m2", "m3", "m4"]

for g in st.session_state.groups:
    origin = g["origin"]
    dist = calculate_haversine(origin, destination)
    g_key = f"{g['name']}_{origin}_{destination}"
    
    current_modes = []
    for idx_m in range(4):
        m_key = default_keys[idx_m]
        state_key = f"name_{g_key}_{m_key}"
        if state_key not in st.session_state.mode_names_state:
            st.session_state.mode_names_state[state_key] = DEFAULT_SPECS[idx_m]["default_name"]
        current_modes.append(st.session_state.mode_names_state[state_key])
    
    with st.expander(f"⚙️ Tùy chỉnh giá & Số hiệu chuyến cho: **{g['name']}** ({origin} ➔ {destination} | {dist} km)", expanded=True):
        mode_options = ["🤖 Tự động đề xuất tối ưu"] + current_modes
        if g_key not in st.session_state.override_mode: st.session_state.override_mode[g_key] = "🤖 Tự động đề xuất tối ưu"
            
        selected_override = st.selectbox(f"🔄 Chọn lại phương tiện cho **{g['name']}**:", mode_options, index=mode_options.index(st.session_state.override_mode[g_key]) if st.session_state.override_mode[g_key] in mode_options else 0, key=f"select_override_{g_key}")
        st.session_state.override_mode[g_key] = selected_override
        
        st.markdown("---")
        st.markdown("**1. Đổi tên phương tiện, Số hiệu chuyến, Giá vé & Giờ đi:**")
        col_m1, col_m2, col_m3, col_m4 = st.columns(4)
        cols = [col_m1, col_m2, col_m3, col_m4]
        
        for idx_m in range(4):
            m_key = default_keys[idx_m]
            spec = DEFAULT_SPECS[idx_m]
            cfg_key = f"{g_key}_{m_key}"
            state_name_key = f"name_{g_key}_{m_key}"
            state_flight_key = f"flight_{g_key}_{m_key}"
            
            if state_flight_key not in st.session_state.flight_codes: st.session_state.flight_codes[state_flight_key] = "VN123" if idx_m == 0 else ""
            default_cost = 0 if dist == 0 else int(spec["base_cost"] + (dist * spec["cost_per_km"]))
            default_dur = 0.0 if dist == 0 else round((dist / spec["speed"]) + spec["fixed_time"], 1)
            if cfg_key not in st.session_state.custom_configs: st.session_state.custom_configs[cfg_key] = {"cost": default_cost, "duration": default_dur}
                
            with cols[idx_m]:
                new_mode_name = st.text_input(f"Phương tiện {idx_m+1}:", value=st.session_state.mode_names_state[state_name_key], key=f"input_name_{g_key}_{m_key}")
                st.session_state.mode_names_state[state_name_key] = new_mode_name
                
                m_check = new_mode_name.lower()
                if "bay" in m_check or "air" in m_check:
                    label_code = "Mã chuyến bay:"
                    ph_code = "VN123 / VJ456"
                elif any(k in m_check for k in ["thủy", "cano", "ca nô", "tàu", "hỏa", "phà", "train", "boat"]):
                    label_code = "Mã chuyến đi:"
                    ph_code = "SE1 / Tàu HQ..."
                else:
                    label_code = "Mã chuyến (nếu có):"
                    ph_code = "Không bắt buộc"
                    
                new_flight_code = st.text_input(label_code, value=st.session_state.flight_codes[state_flight_key], placeholder=ph_code, key=f"input_flight_{g_key}_{m_key}")
                st.session_state.flight_codes[state_flight_key] = new_flight_code
                new_cost = st.number_input(f"Giá/người (VNĐ):", min_value=0, value=st.session_state.custom_configs[cfg_key]["cost"], step=50000, key=f"cost_{cfg_key}")
                new_dur = st.number_input(f"Thời gian (giờ):", min_value=0.0, value=float(st.session_state.custom_configs[cfg_key]["duration"]), step=0.5, key=f"dur_{cfg_key}")
                st.session_state.custom_configs[cfg_key]["cost"] = new_cost
                st.session_state.custom_configs[cfg_key]["duration"] = new_dur

        st.markdown("---")
        st.markdown("**2. Tùy chỉnh văn bản Lộ trình hành trình chi tiết:**")
        active_mode = selected_override if selected_override != "🤖 Tự động đề xuất tối ưu" else current_modes[0]
        active_code = ""
        for i_m in range(4):
            if current_modes[i_m] == active_mode:
                active_code = st.session_state.flight_codes.get(f"flight_{g_key}_{default_keys[i_m]}", "")
                break
                
        default_itinerary_text = generate_itinerary_steps(origin, destination, active_mode, active_code, dist)
        if g_key not in st.session_state.custom_itinerary: st.session_state.custom_itinerary[g_key] = default_itinerary_text
            
        custom_itin_input = st.text_area(f"Sửa nội dung hành trình cho {g['name']}:", value=st.session_state.custom_itinerary[g_key], height=100, key=f"itin_input_{g_key}")
        st.session_state.custom_itinerary[g_key] = custom_itin_input

# -----------------------------------------------------------------------------
# 4. TÍNH TOÁN & HIỂN THỊ KẾT QUẢ
# -----------------------------------------------------------------------------
def calculate_custom_route(g, destination, idx_m):
    origin = g["origin"]
    dist = calculate_haversine(origin, destination)
    m_key = default_keys[idx_m]
    g_key = f"{g['name']}_{origin}_{destination}"
    cfg_key = f"{g_key}_{m_key}"
    
    mode_name = st.session_state.mode_names_state.get(f"name_{g_key}_{m_key}", DEFAULT_SPECS[idx_m]["default_name"])
    flight_code = st.session_state.flight_codes.get(f"flight_{g_key}_{m_key}", "")
    cost_per_person = st.session_state.custom_configs[cfg_key]["cost"]
    duration = st.session_state.custom_configs[cfg_key]["duration"]
    
    spec = DEFAULT_SPECS[idx_m]
    if dist < spec["min_dist"] and dist > 0: return None
        
    total_cost = cost_per_person * g["people"]
    steps = st.session_state.custom_itinerary.get(g_key, generate_itinerary_steps(origin, destination, mode_name, flight_code, dist))
    
    return {
        "name": g["name"],
        "origin": origin,
        "people": g["people"],
        "mode": mode_name,
        "flight_code": flight_code,
        "distance": dist,
        "duration": duration,
        "cost_per_person": cost_per_person,
        "total_cost": total_cost,
        "steps": steps,
        "idx": idx_m
    }

def optimize_travel(groups, destination, criteria):
    results = []
    total_group_cost, max_duration = 0, 0
    for g in groups:
        origin = g["origin"]
        g_key = f"{g['name']}_{origin}_{destination}"
        override_choice = st.session_state.override_mode.get(g_key, "🤖 Tự động đề xuất tối ưu")
        options = [calculate_custom_route(g, destination, i) for i in range(4)]
        options = [o for o in options if o is not None]
        if not options: continue
        
        if override_choice != "🤖 Tự động đề xuất tối ưu":
            best = next((o for o in options if o["mode"] == override_choice), options[0])
        else:
            best = min(options, key=lambda x: x["duration"] if criteria == "time" else x["total_cost"])
        
        total_group_cost += best["total_cost"]
        max_duration = max(max_duration, best["duration"])
        results.append(best)
    return results, total_group_cost, max_duration

def display_plan(results, plan_title, total_cost, max_dur):
    summary_data = []
    for r in results:
        mode_label = format_mode_transport_label(r['mode'], r['flight_code'])
        summary_data.append({
            "Đoàn": r["name"],
            "Điểm xuất phát": r["origin"],
            "Số người": r["people"],
            "Phương tiện lựa chọn": mode_label,
            "Quãng đường": f"{r['distance']} km",
            "Thời gian": f"{r['duration']} giờ",
            "Chi phí / Người": f"{r['cost_per_person']:,} VNĐ",
            "Tổng chi phí đoàn": f"{r['total_cost']:,} VNĐ"
        })
    
    st.markdown("#### 📊 Bảng tổng hợp:")
    st.dataframe(pd.DataFrame(summary_data), use_container_width=True)
    
    col_w, col_p = st.columns(2)
    with col_w:
        docx_file = generate_docx(results, plan_title, destination, total_cost, max_dur)
        st.download_button(
            label="📄 Tải Báo Cáo Word (.docx)",
            data=docx_file,
            file_name=f"Bao_Cao_Phuong_An_{destination}.docx",
            mime="application/vnd.openxmlformats-officedocument.wordprocessingml.document"
        )
    with col_p:
        pdf_file = generate_pdf(results, plan_title, destination, total_cost, max_dur)
        st.download_button(
            label="📑 Tải Báo Cáo PDF",
            data=pdf_file,
            file_name=f"Bao_Cao_Phuong_An_{destination}.pdf",
            mime="application/pdf"
        )
    
    st.markdown("#### 🗺️ Chi tiết lộ trình hành trình từng đoàn:")
    for r in results:
        mode_label = format_mode_transport_label(r['mode'], r['flight_code'])
        with st.expander(f"📍 **{r['name']}** ({r['origin']} ➔ {destination}) | **{mode_label}** | Thời gian: **{r['duration']} giờ**", expanded=True):
            col1, col2 = st.columns([1, 2])
            with col1:
                st.write(f"• **Số lượng người:** {r['people']} khách")
                st.write(f"• **Phương tiện:** {mode_label}")
                st.write(f"• **Quãng đường:** {r['distance']} km")
                st.write(f"• **Chi phí / người:** {r['cost_per_person']:,} VNĐ")
                st.write(f"• **Tổng chi phí đoàn:** {r['total_cost']:,} VNĐ")
            with col2:
                st.write("**Lộ trình chi tiết:**")
                st.text(r["steps"])

st.markdown("---")
st.subheader("📊 Kết quả Phân tích Tự Động / Tùy Chọn")

if not st.session_state.groups:
    st.warning("⚠️ Vui lòng thêm ít nhất 1 đoàn ở thanh bên trái!")
else:
    tab1, tab2 = st.tabs(["⚡ PHƯƠNG ÁN TỐI ƯU THỜI GIAN", "💰 PHƯƠNG ÁN TỐI ƯU CHI PHÍ"])
    
    with tab1:
        res_time, cost_t, dur_t = optimize_travel(st.session_state.groups, destination, "time")
        col_a, col_b, col_c = st.columns(3)
        col_a.metric("⏱️ Thời gian toàn đoàn đến đủ", f"{dur_t} Giờ")
        col_b.metric("💸 Tổng chi phí toàn đoàn", f"{cost_t:,} VNĐ")
        col_c.metric("👥 Tổng số khách", f"{sum(g['people'] for g in st.session_state.groups)} Người")
        st.markdown("---")
        display_plan(res_time, "Tối ưu thời gian", cost_t, dur_t)

    with tab2:
        res_cost, cost_c, dur_c = optimize_travel(st.session_state.groups, destination, "cost")
        col_a, col_b, col_c = st.columns(3)
        col_a.metric("💸 Tổng chi phí toàn đoàn", f"{cost_c:,} VNĐ", delta=f"-{(cost_t - cost_c):,} VNĐ (Tiết kiệm)")
        col_b.metric("⏱️ Thời gian chờ hội quân", f"{dur_c} Giờ")
        col_c.metric("👥 Chi phí TB/người", f"{int(cost_c / sum(g['people'] for g in st.session_state.groups)):,} VNĐ")
        st.markdown("---")
        display_plan(res_cost, "Tối ưu chi phí", cost_c, dur_c)
